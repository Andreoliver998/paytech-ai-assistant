from __future__ import annotations

import io
import re
import zipfile
from xml.sax.saxutils import escape as _xml_escape
from datetime import datetime
from typing import Any, Dict, List, Tuple


def _split_fenced_code(text: str) -> List[Tuple[str, str]]:
    s = text or ""
    parts: List[Tuple[str, str]] = []
    i = 0
    while True:
        start = s.find("```", i)
        if start == -1:
            tail = s[i:]
            if tail:
                parts.append(("text", tail))
            break
        if start > i:
            parts.append(("text", s[i:start]))
        end = s.find("```", start + 3)
        if end == -1:
            parts.append(("text", s[start:]))
            break
        code_block = s[start + 3 : end]
        if "\n" in code_block:
            first, rest = code_block.split("\n", 1)
            if re.match(r"^[a-zA-Z0-9_+-]{1,20}$", first.strip()):
                code_block = rest
        parts.append(("code", code_block))
        i = end + 3
    return parts


def _conversation_title(conv: Dict[str, Any]) -> str:
    t = (conv.get("title") or "Conversa").strip()
    return t or "Conversa"


def render_conversation_docx_bytes(conv: Dict[str, Any]) -> bytes:
    """
    Render DOCX bytes.

    Prefers python-docx when available, but falls back to a minimal pure-Python DOCX writer
    when python-docx (lxml) is unavailable/broken in the environment.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt
    except Exception:
        return render_conversation_docx_bytes_fallback(conv)

    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(3.0)
    sec.left_margin = Cm(3.0)
    sec.bottom_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)

    title = _conversation_title(conv)
    p = doc.add_paragraph()
    run = p.add_run(f"Conversa – {title}")
    run.bold = True
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)

    meta = doc.add_paragraph()
    meta.paragraph_format.first_line_indent = Cm(0)
    meta.add_run(f"Data/hora de exportação: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    meta.add_run(f"Sessão: {conv.get('id') or ''}")
    doc.add_paragraph("")

    for m in conv.get("messages") or []:
        role = (m.get("role") or "").strip().lower()
        label = "Você" if role == "user" else ("Assistente" if role == "assistant" else "Sistema")

        head = doc.add_paragraph()
        head.paragraph_format.first_line_indent = Cm(0)
        r = head.add_run(f"{label}:")
        r.bold = True

        text = m.get("content") or ""
        for kind, part in _split_fenced_code(text):
            if kind == "code":
                code = part.replace("\r\n", "\n").rstrip("\n")
                for line in code.split("\n"):
                    para = doc.add_paragraph()
                    para.paragraph_format.first_line_indent = Cm(0)
                    run = para.add_run(line)
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)
            else:
                blocks = part.replace("\r\n", "\n").split("\n")
                for i, line in enumerate(blocks):
                    para = doc.add_paragraph(line)
                    para.paragraph_format.first_line_indent = Cm(1.25) if i == 0 else Cm(0)

        doc.add_paragraph("")

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def render_conversation_pdf_bytes(conv: Dict[str, Any]) -> bytes:
    """
    Render PDF bytes.

    Prefers reportlab, but falls back to fpdf (pure-Python) when reportlab or its binary deps
    (e.g. pillow) are unavailable/broken.
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle, Preformatted
    except Exception:
        return render_conversation_pdf_bytes_fallback(conv)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=3 * cm,
        rightMargin=2 * cm,
        topMargin=3 * cm,
        bottomMargin=2 * cm,
        title=_conversation_title(conv),
    )

    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "ABNTBase",
        parent=styles["Normal"],
        fontName="Times-Roman",
        fontSize=12,
        leading=12 * 1.5,
        alignment=TA_JUSTIFY,
        firstLineIndent=1.25 * cm,
        spaceAfter=0,
    )
    title_style = ParagraphStyle(
        "ABNTTitle",
        parent=base,
        fontName="Times-Bold",
        fontSize=14,
        leading=14 * 1.5,
        alignment=TA_CENTER,
        firstLineIndent=0,
        spaceAfter=12,
    )
    meta_style = ParagraphStyle(
        "ABNTMeta",
        parent=base,
        alignment=TA_LEFT,
        firstLineIndent=0,
        spaceAfter=10,
    )
    label_style = ParagraphStyle(
        "ABNTLabel",
        parent=base,
        fontName="Times-Bold",
        firstLineIndent=0,
        spaceBefore=10,
        spaceAfter=4,
    )
    code_style = ParagraphStyle(
        "ABNTCode",
        parent=base,
        fontName="Courier",
        fontSize=10,
        leading=10 * 1.2,
        alignment=TA_LEFT,
        firstLineIndent=0,
    )

    def esc(s: str) -> str:
        return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    story = []
    story.append(Paragraph(esc(f"Conversa – {_conversation_title(conv)}"), title_style))
    story.append(Paragraph(esc(f"Data/hora de exportação: {datetime.now().strftime('%Y-%m-%d %H:%M')}"), meta_style))
    story.append(Paragraph(esc(f"Sessão: {conv.get('id') or ''}"), meta_style))
    story.append(Spacer(1, 8))

    sep = Table([[""]], colWidths=[doc.width])
    sep.setStyle(TableStyle([("LINEBELOW", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1"))]))
    story.append(sep)

    for m in conv.get("messages") or []:
        role = (m.get("role") or "").strip().lower()
        label = "Você" if role == "user" else ("Assistente" if role == "assistant" else "Sistema")
        story.append(Spacer(1, 10))
        story.append(Paragraph(esc(f"{label}:"), label_style))

        content = m.get("content") or ""
        for kind, part in _split_fenced_code(content):
            if kind == "code":
                code = part.replace("\r\n", "\n").rstrip("\n")
                pre = Preformatted(esc(code), code_style)
                box = Table([[pre]], colWidths=[doc.width - 0.2 * cm])
                box.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#EEF2FF")),
                            ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#C7D2FE")),
                            ("LEFTPADDING", (0, 0), (-1, -1), 6),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                            ("TOPPADDING", (0, 0), (-1, -1), 6),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                        ]
                    )
                )
                story.append(Spacer(1, 6))
                story.append(box)
                story.append(Spacer(1, 6))
            else:
                text = part.replace("\r\n", "\n")
                paragraphs = text.split("\n\n")
                for para in paragraphs:
                    p = para.strip("\n")
                    if not p:
                        story.append(Spacer(1, 6))
                        continue
                    p = esc(p).replace("\n", "<br/>")
                    story.append(Paragraph(p, base))

        story.append(Spacer(1, 6))
        story.append(sep)

    doc.build(story)
    return buf.getvalue()


def _docx_minimal_package(document_xml: str, title: str) -> bytes:
    """
    Build a minimal DOCX package as a zip of Office Open XML parts.
    Enough for Word/LibreOffice to open as a document.
    """
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>
  <Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>
</Types>
"""
    rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>
  <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>
</Relationships>
"""
    doc_rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>
"""
    now = datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
    core = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
 xmlns:dc="http://purl.org/dc/elements/1.1/"
 xmlns:dcterms="http://purl.org/dc/terms/"
 xmlns:dcmitype="http://purl.org/dc/dcmitype/"
 xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
  <dc:title>{_xml_escape(title or "Conversa")}</dc:title>
  <dc:creator>PayTech AI</dc:creator>
  <cp:lastModifiedBy>PayTech AI</cp:lastModifiedBy>
  <dcterms:created xsi:type="dcterms:W3CDTF">{now}</dcterms:created>
  <dcterms:modified xsi:type="dcterms:W3CDTF">{now}</dcterms:modified>
</cp:coreProperties>
"""
    app = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
 xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
  <Application>PayTech AI</Application>
</Properties>
"""

    bio = io.BytesIO()
    with zipfile.ZipFile(bio, mode="w", compression=zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", content_types)
        z.writestr("_rels/.rels", rels)
        z.writestr("word/document.xml", document_xml)
        z.writestr("word/_rels/document.xml.rels", doc_rels)
        z.writestr("docProps/core.xml", core)
        z.writestr("docProps/app.xml", app)
    return bio.getvalue()


def render_conversation_docx_bytes_fallback(conv: Dict[str, Any]) -> bytes:
    """
    Pure-Python DOCX export. Minimal formatting, but opens reliably without python-docx/lxml.
    """
    title = _conversation_title(conv)
    exported_at = datetime.now().strftime("%Y-%m-%d %H:%M")
    sid = str(conv.get("id") or "")

    def w_p(text: str, bold: bool = False, center: bool = False) -> str:
        t = _xml_escape(text or "")
        jc = '<w:jc w:val="center"/>' if center else ''
        b = "<w:b/>" if bold else ""
        # Preserve leading/trailing spaces
        return (
            f"<w:p><w:pPr>{jc}</w:pPr>"
            f"<w:r><w:rPr>{b}</w:rPr><w:t xml:space=\"preserve\">{t}</w:t></w:r>"
            f"</w:p>"
        )

    paras: List[str] = []
    paras.append(w_p(f"Conversa - {title}", bold=True, center=True))
    paras.append(w_p(f"Data/hora de exportação: {exported_at}"))
    if sid:
        paras.append(w_p(f"Sessão: {sid}"))
    paras.append(w_p(""))

    for m in conv.get("messages") or []:
        role = (m.get("role") or "").strip().lower()
        label = "Você" if role == "user" else ("Assistente" if role == "assistant" else "Sistema")
        paras.append(w_p(f"{label}:", bold=True))
        content = (m.get("content") or "").replace("\r\n", "\n")
        for line in content.split("\n"):
            paras.append(w_p(line))
        paras.append(w_p(""))

    document_xml = (
        """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>"""
        """<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">"""
        """<w:body>"""
        + "".join(paras)
        + """<w:sectPr><w:pgSz w:w="11906" w:h="16838"/><w:pgMar w:top="1701" w:right="1134" w:bottom="1134" w:left="1701"/></w:sectPr>"""
        + """</w:body></w:document>"""
    )
    return _docx_minimal_package(document_xml=document_xml, title=title)


def _pdf_sanitize(text: str) -> str:
    # fpdf (v1.x) is latin-1 oriented; normalize common punctuation outside latin-1.
    s = (text or "")
    s = s.replace("–", "-").replace("—", "-").replace("“", "\"").replace("”", "\"").replace("’", "'").replace("…", "...")
    # ensure only latin-1 characters remain
    return s.encode("latin-1", errors="replace").decode("latin-1", errors="replace")


def render_conversation_pdf_bytes_fallback(conv: Dict[str, Any]) -> bytes:
    """
    Pure-Python PDF export using fpdf (already in requirements). ABNT-ish margins, simple layout.
    """
    from fpdf import FPDF

    title = _conversation_title(conv)
    exported_at = datetime.now().strftime("%Y-%m-%d %H:%M")
    sid = str(conv.get("id") or "")

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=20)
    # ABNT-ish margins: left 30mm, top 30mm, right 20mm
    pdf.set_margins(30, 30, 20)
    pdf.add_page()
    pdf.set_font("Times", "B", 14)
    pdf.multi_cell(0, 8, _pdf_sanitize(f"Conversa - {title}"), align="C")
    pdf.ln(2)
    pdf.set_font("Times", "", 12)
    pdf.multi_cell(0, 7, _pdf_sanitize(f"Data/hora de exportação: {exported_at}"))
    if sid:
        pdf.multi_cell(0, 7, _pdf_sanitize(f"Sessão: {sid}"))
    pdf.ln(2)

    for m in conv.get("messages") or []:
        role = (m.get("role") or "").strip().lower()
        label = "Você" if role == "user" else ("Assistente" if role == "assistant" else "Sistema")
        pdf.ln(2)
        pdf.set_font("Times", "B", 12)
        pdf.multi_cell(0, 7, _pdf_sanitize(f"{label}:"))
        pdf.set_font("Times", "", 12)
        content = (m.get("content") or "").replace("\r\n", "\n")
        pdf.multi_cell(0, 7, _pdf_sanitize(content))
        pdf.ln(1)

    # fpdf v1 returns a latin-1 str for dest='S'
    out = pdf.output(dest="S")
    if isinstance(out, bytes):
        return out
    return str(out).encode("latin-1", errors="replace")
