from __future__ import annotations

import io
import re
from datetime import datetime
from typing import List, Optional, Tuple

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel
from sqlalchemy.orm import Session

from db import get_db
from models import MessageDB, SessionDB


router = APIRouter(prefix="/export", tags=["export"])


class ConversationMessage(BaseModel):
    role: str
    content: str = ""
    ts: Optional[str] = None


class ConversationModel(BaseModel):
    id: str
    title: str = "Conversa"
    createdAt: Optional[str] = None
    updatedAt: Optional[str] = None
    messages: List[ConversationMessage] = []


class ConversationExportRequest(BaseModel):
    conversationId: Optional[str] = None
    conversation: Optional[ConversationModel] = None


def _load_conversation_from_db(db: Session, conversation_id: str) -> ConversationModel:
    sid = (conversation_id or "").strip()
    if not sid:
        raise HTTPException(status_code=400, detail="conversationId inválido.")

    s = db.get(SessionDB, sid)
    if not s:
        raise HTTPException(status_code=404, detail="Conversa não encontrada.")

    msgs = (
        db.query(MessageDB)
        .filter(MessageDB.session_id == sid)
        .order_by(MessageDB.createdAt.asc(), MessageDB.id.asc())
        .all()
    )
    return ConversationModel(
        id=s.id,
        title=s.title or "Conversa",
        createdAt=s.createdAt.isoformat(timespec="seconds") if s.createdAt else None,
        updatedAt=s.updatedAt.isoformat(timespec="seconds") if s.updatedAt else None,
        messages=[
            ConversationMessage(
                role=m.role,
                content=m.content or "",
                ts=m.createdAt.isoformat(timespec="seconds") if m.createdAt else None,
            )
            for m in msgs
        ],
    )


def _slugify(s: str) -> str:
    t = (s or "").strip().lower()
    t = re.sub(r"[^\w\s-]+", "", t, flags=re.UNICODE)
    t = re.sub(r"[\s_-]+", "-", t).strip("-")
    return t or "conversa"


def _export_filename(conv: ConversationModel, ext: str) -> str:
    slug = _slugify(conv.title or "conversa")
    ts = datetime.now().strftime("%Y%m%d-%H%M")
    return f"conversa-{slug}-{ts}.{ext}"


@router.post("/conversation/docx")
def export_conversation_docx(payload: ConversationExportRequest, db: Session = Depends(get_db)):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"python-docx não disponível: {e}")

    conv = payload.conversation or (_load_conversation_from_db(db, payload.conversationId or "") if payload.conversationId else None)
    if conv is None:
        raise HTTPException(status_code=400, detail="Envie `conversationId` ou `conversation`.")

    doc = Document()

    # ABNT-like page setup
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(3.0)
    sec.left_margin = Cm(3.0)
    sec.bottom_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)

    # Title
    p = doc.add_paragraph()
    run = p.add_run(f"Conversa – {conv.title or 'Conversa'}")
    run.bold = True
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)

    meta = doc.add_paragraph()
    meta.paragraph_format.first_line_indent = Cm(0)
    meta.add_run(f"Data/hora de exportação: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    meta.add_run(f"Sessão: {conv.id}")

    doc.add_paragraph("")  # spacer

    for m in conv.messages or []:
        role = (m.role or "").strip().lower()
        label = "Você" if role == "user" else ("Assistente" if role == "assistant" else "Sistema")

        head = doc.add_paragraph()
        head.paragraph_format.first_line_indent = Cm(0)
        r = head.add_run(f"{label}:")
        r.bold = True

        text = m.content or ""
        for kind, part in _split_fenced_code(text):
            if kind == "code":
                code = part.replace("\r\n", "\n").rstrip("\n")
                # bloco de código: monoespaçado, menor, sem recuo ABNT
                for line in code.split("\n"):
                    para = doc.add_paragraph()
                    para.paragraph_format.first_line_indent = Cm(0)
                    run = para.add_run(line)
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)
            else:
                # Texto: preserva quebras de linha
                blocks = part.replace("\r\n", "\n").split("\n")
                for i, line in enumerate(blocks):
                    para = doc.add_paragraph(line)
                    if i == 0:
                        para.paragraph_format.first_line_indent = Cm(1.25)
                    else:
                        para.paragraph_format.first_line_indent = Cm(0)

        doc.add_paragraph("")  # spacer

    bio = io.BytesIO()
    doc.save(bio)
    data = bio.getvalue()

    filename = _export_filename(conv, "docx")
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _split_fenced_code(text: str) -> List[Tuple[str, str]]:
    """
    Returns list of (kind, content) where kind is 'text' or 'code'.
    Preserves code content exactly between ``` fences.
    """
    s = text or ""
    parts: List[Tuple[str, str]] = []
    i = 0
    while True:
        start = s.find("```", i)
        if start == -1:
            tail = s[i:]
            if tail:
                parts.append(("text", tail))
            break
        if start > i:
            parts.append(("text", s[i:start]))
        end = s.find("```", start + 3)
        if end == -1:
            # no closing fence -> treat rest as text (rule: don't invent)
            parts.append(("text", s[start:]))
            break
        code_block = s[start + 3 : end]
        # drop optional language first line, keep rest
        if "\n" in code_block:
            first, rest = code_block.split("\n", 1)
            if re.match(r"^[a-zA-Z0-9_+-]{1,20}$", first.strip()):
                code_block = rest
        parts.append(("code", code_block))
        i = end + 3
    return parts


@router.post("/conversation/pdf")
def export_conversation_pdf(payload: ConversationExportRequest, db: Session = Depends(get_db)):
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle, Preformatted
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"reportlab não disponível: {e}")

    conv = payload.conversation or (_load_conversation_from_db(db, payload.conversationId or "") if payload.conversationId else None)
    if conv is None:
        raise HTTPException(status_code=400, detail="Envie `conversationId` ou `conversation`.")
    filename = _export_filename(conv, "pdf")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=3 * cm,
        rightMargin=2 * cm,
        topMargin=3 * cm,
        bottomMargin=2 * cm,
        title=conv.title or "Conversa",
    )

    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "ABNTBase",
        parent=styles["Normal"],
        fontName="Times-Roman",
        fontSize=12,
        leading=12 * 1.5,
        alignment=TA_JUSTIFY,
        firstLineIndent=1.25 * cm,
        spaceAfter=0,
    )
    title_style = ParagraphStyle(
        "ABNTTitle",
        parent=base,
        fontName="Times-Bold",
        fontSize=14,
        leading=14 * 1.5,
        alignment=TA_CENTER,
        firstLineIndent=0,
        spaceAfter=12,
    )
    meta_style = ParagraphStyle(
        "ABNTMeta",
        parent=base,
        alignment=TA_LEFT,
        firstLineIndent=0,
        spaceAfter=10,
    )
    label_style = ParagraphStyle(
        "ABNTLabel",
        parent=base,
        fontName="Times-Bold",
        firstLineIndent=0,
        spaceBefore=10,
        spaceAfter=4,
    )
    code_style = ParagraphStyle(
        "ABNTCode",
        parent=base,
        fontName="Courier",
        fontSize=10,
        leading=10 * 1.2,
        alignment=TA_LEFT,
        firstLineIndent=0,
    )

    def esc(s: str) -> str:
        return (
            (s or "")
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

    story = []
    story.append(Paragraph(esc(f"Conversa – {conv.title or 'Conversa'}"), title_style))
    story.append(Paragraph(esc(f"Data/hora de exportação: {datetime.now().strftime('%Y-%m-%d %H:%M')}"), meta_style))
    story.append(Paragraph(esc(f"Sessão: {conv.id}"), meta_style))
    story.append(Spacer(1, 8))

    sep = Table([[""]], colWidths=[doc.width])
    sep.setStyle(TableStyle([("LINEBELOW", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1"))]))
    story.append(sep)

    for m in conv.messages or []:
        role = (m.role or "").strip().lower()
        label = "Você" if role == "user" else ("Assistente" if role == "assistant" else "Sistema")
        story.append(Spacer(1, 10))
        story.append(Paragraph(esc(f"{label}:"), label_style))

        content = m.content or ""
        for kind, part in _split_fenced_code(content):
            if kind == "code":
                code = part.replace("\r\n", "\n").rstrip("\n")
                pre = Preformatted(esc(code), code_style)
                box = Table([[pre]], colWidths=[doc.width - 0.2 * cm])
                box.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#EEF2FF")),
                            ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#C7D2FE")),
                            ("LEFTPADDING", (0, 0), (-1, -1), 6),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                            ("TOPPADDING", (0, 0), (-1, -1), 6),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                        ]
                    )
                )
                story.append(Spacer(1, 6))
                story.append(box)
                story.append(Spacer(1, 6))
            else:
                text = part.replace("\r\n", "\n")
                paragraphs = text.split("\n\n")
                for para in paragraphs:
                    p = para.strip("\n")
                    if not p:
                        story.append(Spacer(1, 6))
                        continue
                    # preserve single newlines inside paragraph
                    p = esc(p).replace("\n", "<br/>")
                    story.append(Paragraph(p, base))

        story.append(Spacer(1, 6))
        story.append(sep)

    doc.build(story)
    data = buf.getvalue()

    return Response(
        content=data,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
